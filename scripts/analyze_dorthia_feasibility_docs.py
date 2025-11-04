"""
Analyze Dorthia's feasibility report documents and extract key feedback
"""

from docx import Document
import os
import json

def extract_text_from_docx(filepath):
    """Extract all text from a Word document"""
    try:
        doc = Document(filepath)
        full_text = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(f"[TABLE] {cell.text}")

        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading {filepath}: {str(e)}"

def analyze_documents():
    """Analyze all three feasibility report documents"""

    docs = [
        r"C:\Users\MarieLexisDad\Downloads\NURS 8123 Feasibility Report - Dorthia Daudier.docx",
        r"C:\Users\MarieLexisDad\Downloads\DDaudier - Feasibility Revised.docx",
        r"C:\Users\MarieLexisDad\Downloads\Dorthia Daudier Feasibility Draft (1).docx"
    ]

    results = {}

    for doc_path in docs:
        filename = os.path.basename(doc_path)
        print(f"\n{'='*80}")
        print(f"ANALYZING: {filename}")
        print('='*80)

        text = extract_text_from_docx(doc_path)
        results[filename] = text

        # Identify key sections (don't print to avoid encoding issues)
        sections_to_find = [
            "PICOT",
            "Problem Statement",
            "Purpose",
            "Feasibility",
            "Measurement",
            "Evaluation",
            "Timeline",
            "Budget",
            "Evidence",
            "Framework",
            "Self-Efficacy",
            "First Call Resolution",
            "FCR",
            "Feedback",
            "Revision",
            "Comment"
        ]

        found_sections = []
        for section in sections_to_find:
            if section.lower() in text.lower():
                found_sections.append(section)

        results[filename + "_sections"] = found_sections

    # Save full extracted text
    output_file = "scripts/dorthia_feasibility_analysis.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)

    print(f"\n\n{'='*80}")
    print(f"Full analysis saved to: {output_file}")
    print('='*80)

if __name__ == "__main__":
    analyze_documents()
